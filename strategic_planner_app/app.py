import io
import re
from datetime import datetime
from typing import Dict, List, Tuple, Any, Optional

import streamlit as st
from docx import Document
import requests


# ---------------------------
# Burke–Litwin Model Keywords
# ---------------------------
BURKE_LITWIN_FACTORS: Dict[str, Dict[str, Any]] = {
    "External Environment": {
        "keywords": [
            "market", "regulation", "regulatory", "competition", "competitor", "technology", "economic",
            "customer", "client", "industry", "trend", "political", "legal", "supply chain", "inflation",
            "disruption", "pandemic", "geopolitical", "vendor", "partner", "ecosystem", "benchmark"
        ]
    },
    "Mission & Strategy": {
        "keywords": [
            "mission", "vision", "strategy", "strategic", "north star", "roadmap", "goal", "objectives",
            "priorities", "value proposition", "positioning", "differentiation", "portfolio", "investment",
            "okrs", "kpi", "target state", "business model"
        ]
    },
    "Leadership": {
        "keywords": [
            "leadership", "executive", "management", "sponsor", "steerco", "steering committee", "decision",
            "visionary", "accountability", "role modeling", "tone from the top", "buy-in", "support"
        ]
    },
    "Organizational Culture": {
        "keywords": [
            "culture", "values", "norms", "behaviors", "mindset", "trust", "collaboration", "innovation",
            "risk appetite", "psychological safety", "inclusion", "ownership", "agile", "customer-centric"
        ]
    },
    "Structure": {
        "keywords": [
            "structure", "org design", "organizational design", "hierarchy", "matrix", "span of control",
            "reporting", "governance", "roles", "responsibilities", "raci", "centralized", "decentralized",
            "operating model"
        ]
    },
    "Systems (Policies & Processes)": {
        "keywords": [
            "process", "policy", "procedure", "workflow", "standard", "sop", "tooling", "systems",
            "automation", "controls", "compliance", "documentation", "change control", "itil", "qms"
        ]
    },
    "Management Practices": {
        "keywords": [
            "performance management", "cadence", "reviews", "1:1", "one-on-one", "coaching", "mentoring",
            "planning", "prioritization", "resource allocation", "budgeting", "forecasting", "reporting",
            "status", "standup", "retrospective"
        ]
    },
    "Work Unit Climate": {
        "keywords": [
            "morale", "engagement", "turnover", "conflict", "communication", "collaboration", "silos",
            "workload", "burnout", "stress", "psychological safety", "recognition", "feedback", "trust"
        ]
    },
    "Task & Individual Skills": {
        "keywords": [
            "skill", "capability", "competency", "training", "enablement", "expertise", "upskilling",
            "reskilling", "certification", "knowledge", "experience", "talent", "hiring", "recruiting"
        ]
    },
    "Individual Needs & Values": {
        "keywords": [
            "motivation", "purpose", "recognition", "reward", "benefits", "wellbeing", "welfare", "career",
            "growth", "development", "flexibility", "work-life", "inclusion", "belonging"
        ]
    },
    "Motivation": {
        "keywords": [
            "engagement", "commitment", "initiative", "ownership", "empower", "autonomy", "incentive",
            "bonus", "compensation", "rewards", "celebrate", "acknowledge"
        ]
    },
    "Performance (Individual & Org)": {
        "keywords": [
            "kpi", "okr", "metric", "sla", "quality", "throughput", "cycle time", "lead time",
            "revenue", "profit", "margin", "nps", "customer satisfaction", "error rate", "defect",
            "utilization", "productivity", "churn"
        ]
    },
}


DEFAULT_CAUSAL_LINKS: List[Tuple[str, str]] = [
    ("External Environment", "Mission & Strategy"),
    ("Leadership", "Organizational Culture"),
    ("Leadership", "Management Practices"),
    ("Organizational Culture", "Work Unit Climate"),
    ("Mission & Strategy", "Structure"),
    ("Structure", "Systems (Policies & Processes)"),
    ("Systems (Policies & Processes)", "Management Practices"),
    ("Management Practices", "Work Unit Climate"),
    ("Task & Individual Skills", "Motivation"),
    ("Individual Needs & Values", "Motivation"),
    ("Work Unit Climate", "Motivation"),
    ("Motivation", "Performance (Individual & Org)"),
]


def _normalize_text(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[\r\n]+", " \n ", text)
    text = re.sub(r"[^a-z0-9\s\-\&]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_text_from_docx(uploaded) -> str:
    document = Document(uploaded)
    paragraphs: List[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            paragraphs.append(paragraph.text.strip())
    # Extract text from simple tables as well
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)
    return "\n".join(paragraphs)


def analyze_text_burke_litwin(text: str) -> Dict[str, Any]:
    normalized = _normalize_text(text)
    factor_results: Dict[str, Any] = {}

    for factor, meta in BURKE_LITWIN_FACTORS.items():
        keywords = meta.get("keywords", [])
        signals: List[Tuple[str, int]] = []
        total_hits = 0
        for kw in keywords:
            pattern = r"\b" + re.escape(kw.lower()) + r"\b"
            matches = re.findall(pattern, normalized)
            if matches:
                count = len(matches)
                total_hits += count
                signals.append((kw, count))
        # Heuristic: more mentions -> higher salience. Convert to score 0-100 with soft cap.
        score = min(100, int(10 * (total_hits ** 0.5) * 3)) if total_hits > 0 else 25
        factor_results[factor] = {
            "mentions": total_hits,
            "score": score,
            "signals": sorted(signals, key=lambda x: x[1], reverse=True),
        }

    # Identify simple causal priorities: upstream factors with high mentions and low related downstream scores
    causal_links: List[Dict[str, Any]] = []
    for source, target in DEFAULT_CAUSAL_LINKS:
        s = factor_results[source]["score"]
        t = factor_results[target]["score"]
        weight = max(1, int((s / 10) - (t / 20)))
        if weight > 0:
            causal_links.append({"from": source, "to": target, "weight": weight})

    # Prioritize areas needing improvement: low score or low mentions
    priorities = sorted(
        (
            (factor, data["score"], data["mentions"]) for factor, data in factor_results.items()
        ),
        key=lambda x: (x[1], x[2])
    )

    return {
        "factors": factor_results,
        "links": causal_links,
        "priorities": priorities,
    }


def _owner_by_factor(factor: str) -> str:
    mapping = {
        "External Environment": "Strategy Office",
        "Mission & Strategy": "CEO / Strategy Office",
        "Leadership": "Executive Leadership Team",
        "Organizational Culture": "HR / People",
        "Structure": "COO / Org Design",
        "Systems (Policies & Processes)": "Operations / Compliance",
        "Management Practices": "People Managers",
        "Work Unit Climate": "HRBPs / Department Heads",
        "Task & Individual Skills": "Learning & Development",
        "Individual Needs & Values": "HR / Total Rewards",
        "Motivation": "HR / People Managers",
        "Performance (Individual & Org)": "CFO / COO",
    }
    return mapping.get(factor, "Executive Sponsor")


def generate_strategic_plan(analysis: Dict[str, Any], horizon_months: int = 12, max_objectives: int = 6) -> Dict[str, Any]:
    # Select top needs (lowest scores)
    priorities_sorted = analysis["priorities"][:max_objectives]

    objectives: List[Dict[str, Any]] = []
    initiatives: List[Dict[str, Any]] = []

    for factor, score, mentions in priorities_sorted:
        objective = {
            "title": f"Improve {factor} maturity",
            "description": f"Address gaps in {factor.lower()} revealed by the current-state document to enable performance uplift.",
            "target": "Raise factor score to 70+ and reduce key risks.",
            "owner": _owner_by_factor(factor),
            "kpis": [
                f"{factor} composite score",
                "On-time milestone delivery",
                "Stakeholder satisfaction index",
            ],
        }
        objectives.append(objective)

        # Transactional and transformational initiatives per factor
        initiatives.extend([
            {
                "factor": factor,
                "type": "Transformational",
                "title": f"Design and implement target-state for {factor}",
                "owner": _owner_by_factor(factor),
                "timeline": "Q1–Q4",
                "kpis": ["Design sign-off", "Adoption rate", "Outcome metrics trending"],
                "risks": ["Change resistance", "Under-resourcing", "Scope creep"],
                "dependencies": ["Executive sponsorship", "Change management", "Budget approval"],
            },
            {
                "factor": factor,
                "type": "Transactional",
                "title": f"Quick wins to stabilize {factor.lower()}",
                "owner": _owner_by_factor(factor),
                "timeline": "Q1–Q2",
                "kpis": ["Cycle time reduction", "Defect rate", "SLA adherence"],
                "risks": ["Competing priorities"],
                "dependencies": ["Process owners", "Tooling access"],
            },
        ])

    plan = {
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
        "horizon_months": horizon_months,
        "objectives": objectives,
        "initiatives": initiatives,
        "analysis": analysis,
    }
    return plan


def build_plan_docx(plan: Dict[str, Any], source_title: str = "Uploaded Document", llm_narrative: Optional[str] = None) -> bytes:
    doc = Document()
    doc.add_heading('Strategic Plan based on Burke–Litwin Model', 0)
    doc.add_paragraph(f"Generated: {plan['generated_at']}")
    doc.add_paragraph(f"Source: {source_title}")

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This strategic plan translates observed organizational signals into actionable objectives and initiatives "
        "using the Burke–Litwin causal model to drive sustainable performance outcomes."
    )

    doc.add_heading('Diagnosis by Burke–Litwin Factor', level=1)
    for factor, data in plan["analysis"]["factors"].items():
        doc.add_heading(factor, level=2)
        doc.add_paragraph(f"Score: {data['score']} | Mentions: {data['mentions']}")
        if data["signals"]:
            doc.add_paragraph("Top signals:")
            for kw, count in data["signals"][:5]:
                doc.add_paragraph(f"{kw} ({count})", style='List Bullet')

    doc.add_heading('Strategic Objectives', level=1)
    for idx, obj in enumerate(plan["objectives"], start=1):
        doc.add_heading(f"O{idx}. {obj['title']}", level=2)
        doc.add_paragraph(obj["description"])
        doc.add_paragraph(f"Target: {obj['target']}")
        doc.add_paragraph(f"Owner: {obj['owner']}")
        doc.add_paragraph("KPIs:")
        for k in obj["kpis"]:
            doc.add_paragraph(k, style='List Bullet')

    doc.add_heading('Initiatives & Roadmap', level=1)
    for init in plan["initiatives"]:
        doc.add_heading(f"{init['type']}: {init['title']}", level=2)
        doc.add_paragraph(f"Factor: {init['factor']}")
        doc.add_paragraph(f"Owner: {init['owner']}")
        doc.add_paragraph(f"Timeline: {init['timeline']}")
        doc.add_paragraph("KPIs:")
        for k in init["kpis"]:
            doc.add_paragraph(k, style='List Bullet')
        doc.add_paragraph("Risks:")
        for r in init["risks"]:
            doc.add_paragraph(r, style='List Bullet')
        doc.add_paragraph("Dependencies:")
        for d in init["dependencies"]:
            doc.add_paragraph(d, style='List Bullet')

    doc.add_heading('Causal Links Considered', level=1)
    for link in plan["analysis"]["links"]:
        doc.add_paragraph(f"{link['from']} → {link['to']} (weight {link['weight']})")

    if llm_narrative:
        doc.add_heading('LLM-Refined Strategic Plan (Ollama LLaMA 3.1)', level=1)
        for line in llm_narrative.splitlines():
            text = line.strip('\n')
            if not text:
                doc.add_paragraph("")
            else:
                doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _analysis_table(analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for factor, data in analysis["factors"].items():
        rows.append({
            "Factor": factor,
            "Score": data["score"],
            "Mentions": data["mentions"],
            "Top signals": ", ".join([f"{kw}({cnt})" for kw, cnt in data["signals"][:3]])
        })
    rows.sort(key=lambda r: r["Score"])  # show lowest first
    return rows


def _build_llm_prompts(raw_text: str, analysis: Dict[str, Any]) -> Tuple[str, str]:
    table_rows = _analysis_table(analysis)
    table_text_lines: List[str] = [
        "Factor | Score | Mentions | Top signals"
    ]
    for r in table_rows:
        table_text_lines.append(f"{r['Factor']} | {r['Score']} | {r['Mentions']} | {r['Top signals']}")
    table_text = "\n".join(table_text_lines)

    truncated_text = raw_text[:6000]

    system_prompt = (
        "You are an elite management consultant specializing in organizational transformation using the Burke–Litwin model. "
        "Synthesize inputs into a pragmatic, board-ready strategic plan. Be specific, concise, and actionable."
    )

    user_prompt = (
        "Context document (truncated):\n" + truncated_text +
        "\n\nObserved factor signals (heuristic):\n" + table_text +
        "\n\nDeliver a Burke–Litwin-based strategic plan with these sections: \n"
        "1) Executive Summary (2-4 bullets)\n"
        "2) Diagnosis by Factor (insights per factor)\n"
        "3) Strategic Objectives (3-8, each with owner, target, KPIs)\n"
        "4) Initiatives (transformational and quick wins, with timeline, risks, dependencies)\n"
        "5) KPIs Dashboard (bulleted)\n"
        "6) Causal Links and Rationale (bulleted)\n"
        "Keep it structured in clear markdown."
    )
    return system_prompt, user_prompt


def call_ollama_chat(host: str, port: int, model: str, system_prompt: str, user_prompt: str, temperature: float = 0.3, timeout_s: int = 120) -> str:
    url = f"http://{host}:{port}/api/chat"
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "stream": False,
        "options": {"temperature": max(0.0, min(1.0, temperature))},
    }
    try:
        resp = requests.post(url, json=payload, timeout=timeout_s)
        resp.raise_for_status()
        data = resp.json()
        # Ollama chat returns { message: { content: "..." }, ... }
        message = data.get("message", {})
        content = message.get("content")
        if not content and isinstance(data, dict) and "choices" in data:
            # Fallback if API surface differs
            choices = data.get("choices") or []
            if choices:
                content = choices[0].get("message", {}).get("content")
        return content or ""
    except Exception as e:
        return f"[LLM error] {e}"


def main() -> None:
    st.set_page_config(page_title="Burke–Litwin Strategic Planner", page_icon="📄", layout="wide")
    st.title("Burke–Litwin Strategic Planning (Document-Driven)")
    st.write(
        "Upload a Word document describing your organization, strategy, or current state. "
        "This app will analyze signals across the Burke–Litwin causal model and produce a structured strategic plan "
        "that you can download as a Word document."
    )

    with st.sidebar:
        st.header("Settings")
        horizon = st.slider("Planning horizon (months)", min_value=3, max_value=24, value=12, step=3)
        max_objectives = st.slider("Max strategic objectives", min_value=3, max_value=10, value=6)
        st.divider()
        st.subheader("Ollama (LLM)")
        use_llm = st.checkbox("Use Ollama LLaMA 3.1 to refine plan", value=True)
        ollama_host = st.text_input("Ollama host", value="192.168.2.200")
        ollama_port = st.number_input("Ollama port", value=11434, step=1)
        ollama_model = st.text_input("Model", value="llama3.1")
        ollama_temperature = st.slider("Temperature", min_value=0.0, max_value=1.0, value=0.3, step=0.1)

    uploaded = st.file_uploader("Upload .docx", type=["docx"], accept_multiple_files=False)
    if uploaded is None:
        st.info("Please upload a .docx file to begin.")
        return

    try:
        raw_text = extract_text_from_docx(uploaded)
    except Exception as e:
        st.error(f"Failed to read the document: {e}")
        return

    if not raw_text.strip():
        st.warning("The uploaded document appears to be empty or unreadable.")
        return

    st.subheader("Document Preview")
    st.text_area("Extracted text (first 2,500 chars)", value=raw_text[:2500], height=200)

    if st.button("Generate Strategic Plan", type="primary"):
        with st.spinner("Analyzing document and generating plan..."):
            analysis = analyze_text_burke_litwin(raw_text)
            plan = generate_strategic_plan(analysis, horizon_months=horizon, max_objectives=max_objectives)

        st.success("Plan generated.")

        st.subheader("Analysis Summary (Burke–Litwin)")
        st.dataframe(_analysis_table(analysis), use_container_width=True)

        st.subheader("Proposed Plan (Preview)")
        for idx, obj in enumerate(plan["objectives"], start=1):
            st.markdown(f"**O{idx}. {obj['title']}**")
            st.write(obj["description"])
            st.caption(f"Owner: {obj['owner']} | Target: {obj['target']}")

        llm_output: Optional[str] = None
        if use_llm:
            with st.spinner("Contacting Ollama to refine the plan..."):
                sys_p, usr_p = _build_llm_prompts(raw_text, analysis)
                llm_output = call_ollama_chat(
                    host=ollama_host,
                    port=int(ollama_port),
                    model=ollama_model,
                    system_prompt=sys_p,
                    user_prompt=usr_p,
                    temperature=float(ollama_temperature),
                )
            st.subheader("LLM-Refined Plan Narrative")
            if llm_output and not llm_output.startswith("[LLM error]"):
                st.markdown(llm_output)
            else:
                st.error(llm_output or "LLM returned empty response.")

        # Build downloadable Word doc
        try:
            doc_bytes = build_plan_docx(plan, source_title=uploaded.name, llm_narrative=llm_output)
            st.download_button(
                label="Download Strategic Plan (.docx)",
                data=doc_bytes,
                file_name="Strategic_Plan_Burke_Litwin.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.error(f"Failed to generate Word document: {e}")


if __name__ == "__main__":
    main()

