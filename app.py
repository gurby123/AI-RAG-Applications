import os
import io
import re
from typing import List, Dict

import streamlit as st
import requests
from docx import Document

# Defaults (can be overridden in the UI sidebar or via env vars)
DEFAULT_OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://192.168.2.200:11434")
DEFAULT_OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.1:latest")

FRAMEWORKS = [
    "Reframing Thinking",
    "Delphi",
    "SCAMPER",
    "Blue Ocean",
    "Six Thinking Hats",
    "Balanced Scorecard",
    "McKinsey 7S",
    "Burke-Litwin",
    "TRIZ",
]

FRAMEWORK_DESCRIPTIONS = {
    "Reframing Thinking": "Challenge assumptions and reframe the problem from alternative perspectives.",
    "Delphi": "Iterative expert consensus via rounds of anonymous feedback.",
    "SCAMPER": "Substitute, Combine, Adapt, Modify, Put to another use, Eliminate, Reverse.",
    "Blue Ocean": "Create new market space by shifting value curves away from competition.",
    "Six Thinking Hats": "Parallel thinking: Facts, Feelings, Caution, Benefits, Creativity, Process.",
    "Balanced Scorecard": "Translate strategy into objectives across Financial, Customer, Internal, Learning.",
    "McKinsey 7S": "Align Strategy, Structure, Systems, Shared Values, Skills, Style, Staff.",
    "Burke-Litwin": "Diagnose org change via External Env., Leadership, Culture, Systems, etc.",
    "TRIZ": "Inventive problem solving using patterns of technical evolution and contradictions.",
}

FRAMEWORK_TIPS = {
    "Reframing Thinking": "List assumptions to challenge and alternative frames.",
    "Delphi": "What expert hypotheses or consensus themes might emerge?",
    "SCAMPER": "Ideas for Substitute, Combine, Adapt, Modify, Put to other use, Eliminate, Reverse.",
    "Blue Ocean": "What noncustomers, new value curves, or eliminate-reduce-raise-create actions?",
    "Six Thinking Hats": "Facts, feelings, risks, benefits, creative options, process checks.",
    "Balanced Scorecard": "Objectives across Financial, Customer, Internal Process, Learning & Growth.",
    "McKinsey 7S": "Implications for Strategy, Structure, Systems, Shared Values, Skills, Style, Staff.",
    "Burke-Litwin": "External drivers, leadership, culture, work unit climate, systems alignment.",
    "TRIZ": "Contradictions to resolve, ideality, inventive principles to apply.",
}


def ollama_generate(host: str, model: str, system_prompt: str, user_prompt: str, temperature: float = 0.7) -> str:
    """Call Ollama's /api/generate endpoint for a single-turn completion."""
    url = f"{host.rstrip('/')}/api/generate"
    payload = {
        "model": model,
        "prompt": f"<|system|>\n{system_prompt}\n<|user|>\n{user_prompt}\n<|assistant|>",
        "stream": False,
        "options": {"temperature": temperature},
    }
    resp = requests.post(url, json=payload, timeout=120)
    resp.raise_for_status()
    data = resp.json()
    return data.get("response", "").strip()


def build_vision_prompt(issue_text: str, selected_frameworks: List[str]) -> str:
    frameworks_block = "\n".join(
        f"- {fw}: {FRAMEWORK_DESCRIPTIONS.get(fw, '')}" for fw in selected_frameworks
    )
    return (
        "You are a strategy consultant. From the text below (which may include uploaded document content), first infer 5–10 key concepts/themes, then use those concepts with the chosen frameworks to produce exactly 3 distinct vision statements as a numbered list (1., 2., 3.). Do not print the concepts; only output the three numbered items.\n\n"
        f"Business issues and/or document text:\n{issue_text}\n\n"
        f"Frameworks to apply:\n{frameworks_block}\n\n"
        "Vision statements should be: forward-looking; motivating and inspirational; reflective of culture and core values; aimed at future benefits; and define the organization's destination and reason for existence.\n"
        "Consider: major issues/problems; major strengths/assets; desired changes; purpose of the organization; the kind of organization to create; the dream/vision; hope for a better future; inspiration for effective action; a basis for action planning; and why these issues matter.\n"
        "Make each a positive, declarative one-sentence picture of what the organization will be like in 10–15 years. Focus on themes with long-term direction, activities to pursue, capabilities to build, and purposeful action.\n"
        "Style examples (do not copy): 'To become the world’s most loved, most flown, and most profitable airline.'; 'To create a better everyday life for many people.'; 'Create the most compelling car company of the 21st century by driving the world’s transition to electric vehicles.'\n\n"
        "Instructions:\n"
        "- Each vision statement must be ONE sentence only (<= 25 words).\n"
        "- Ensure each is meaningfully different and grounded in the inferred key concepts.\n"
        "- Output ONLY the three numbered items, no preface or epilogue.\n"
    )


def build_framework_vision_prompt(issue_text: str, framework: str, notes: str) -> str:
    desc = FRAMEWORK_DESCRIPTIONS.get(framework, "")
    tip = FRAMEWORK_TIPS.get(framework, "")
    return (
        f"Using the {framework} framework, generate exactly 3 one-sentence vision statements as a numbered list (1., 2., 3.) based on the text and notes below. Do not print any analysis, only the three items.\n\n"
        f"Framework lens: {framework} — {desc}\n"
        f"Guidance: {tip}\n\n"
        f"Text/context:\n{issue_text}\n\n"
        f"Framework notes:\n{notes}\n\n"
        "Vision statements must be: forward-looking, inspiring, values-aligned, benefit-oriented, clearly defining destination and purpose; ONE sentence (<= 25 words) each; meaningfully distinct."
    )


def parse_numbered_list(text: str) -> List[str]:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    items: List[str] = []
    current: List[str] = []
    for ln in lines:
        if any(ln.startswith(prefix) for prefix in ("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            if current:
                items.append(" ".join(current).strip())
                current = []
            current.append(ln.split(".", 1)[1].strip())
        else:
            current.append(ln)
    if current:
        items.append(" ".join(current).strip())
    return [it for it in items if it]


def first_sentence(text: str) -> str:
    for ch in ".!?":
        pass
    sentence_end = None
    for idx, ch in enumerate(text):
        if ch in ".!?":
            sentence_end = idx
            break
    return text[: sentence_end + 1].strip() if sentence_end is not None else text.strip()


def clean_mission_text(resp_text: str, mission_text: str) -> str:
    cleaned = re.sub(r"^mission(?:\s*statement)?\s*[-:]*\s*", "", mission_text, flags=re.IGNORECASE).strip()
    placeholders = {"mission", "mission statement", "tbd", "placeholder", "n/a", "mission:", "mission statement:"}
    if cleaned.lower() in placeholders or len(cleaned.split()) < 3:
        after = resp_text.split("Mission:", 1)[1] if "Mission:" in resp_text else resp_text
        for line in after.splitlines():
            s = re.sub(r"^mission(?:\s*statement)?\s*[-:]*\s*", "", line.strip(), flags=re.IGNORECASE).strip()
            if not s:
                continue
            if s.lower().startswith("goals"):
                break
            if len(s) >= 2 and s[0].isdigit() and s[1:2] == ".":
                continue
            if len(s.split()) >= 3 and s.lower() not in placeholders:
                return first_sentence(s)
        return cleaned
    return cleaned


def is_meta_mission(text: str) -> bool:
    t = text.strip()
    lower = t.lower()
    if not lower:
        return True
    if lower.startswith("our mission is") or lower.startswith("to "):
        return False
    meta_fragments = [
        "based on the selected vision",
        "i will draft",
        "we will draft",
        "i will develop",
        "we will develop",
        "this mission statement",
        "the mission statement",
        "in this section",
        "i will create",
        "we will create a mission",
        "this will outline",
        "the following mission",
        "here is the mission",
    ]
    if any(frag in lower for frag in meta_fragments):
        return True
    if len(t.split()) < 4:
        return True
    return False


def synthesize_mission_from_vision(vision: str) -> str:
    v = vision.strip().rstrip(".")
    if not v:
        return ""
    if v.lower().startswith("to "):
        rest = v[3:]
        return f"Our mission is to {rest[0].lower()}{rest[1:]} .".replace("  ", " ")
    return f"Our mission is to {v[0].lower()}{v[1:]} .".replace("  ", " ")


def generate_mission_only(host: str, model: str, vision: str, issue_text: str, temperature: float) -> str:
    user_prompt = (
        "From the selected vision below, write ONLY the mission statement as 1–2 sentences. "
        "Do NOT include labels, prefaces, or describe the task. Output the mission content only.\n\n"
        f"Selected vision (one sentence):\n{vision}\n\n"
        f"Context:\n{issue_text}\n"
    )
    resp = ollama_generate(
        host=host,
        model=model,
        system_prompt="Return only the mission statement text, no labels.",
        user_prompt=user_prompt,
        temperature=temperature,
    )
    for line in resp.splitlines():
        s = line.strip()
        if s:
            s = re.sub(r"^mission(?:\s*statement)?\s*[-:]*\s*", "", s, flags=re.IGNORECASE).strip()
            return first_sentence(s) if "." in s or "!" in s or "?" in s else s
    return resp.strip()


def build_mission_goals_prompt(selected_vision: str, issue_text: str, selected_frameworks: List[str]) -> str:
    frameworks_block = ", ".join(selected_frameworks)
    return (
        "You are a strategy consultant. Based on the selected vision statement, draft a single mission statement and at least 5 strategic goals. The mission must expand and operationalize the chosen vision; do NOT use placeholders like 'Mission Statement' or 'TBD'. Do NOT describe what you are about to do; write the mission content itself. Write 'Mission: ' followed immediately by the mission content on the same line. Each goal must clearly align with and advance both the vision and the mission.\n\n"
        f"Selected vision (one sentence):\n{selected_vision}\n\n"
        f"Key business issues and/or document text:\n{issue_text}\n\n"
        f"Framework lenses to reflect: {frameworks_block}.\n\n"
        "Mission guidance:\n"
        "- Describe what the organization will do and why, for whom, and how (the differentiating approach).\n"
        "- Concise (prefer 1 sentence, at most 2), outcome-oriented, inclusive of goals and stakeholders.\n"
        "- Overarching expression of purpose that will drive goals, initiatives, and objectives.\n\n"
        "Goals guidance:\n"
        "- Provide at least 5 higher-order strategic goals aligned to the mission and vision, looking ~5 years out.\n"
        "- Keep goals generic at this stage (not KPIs), ambitious, challenging, inspiring, and avoid a laundry list.\n"
        "- Derive from mission, vision, and capabilities; do not copy from other organizations.\n"
        "- Ensure each goal is traceably aligned to the mission and advances the vision.\n\n"
        "Output format (strict):\n"
        "Mission: <one concise mission statement (1–2 sentences)>\n"
        "Goals:\n"
        "1. <goal one>\n2. <goal two>\n3. <goal three>\n4. <goal four>\n5. <goal five>\n"
    )


def export_docx(issue_text: str, frameworks: List[str], visions: List[str], chosen_idx: int, mission: str, goals: List[str]) -> bytes:
    doc = Document()
    doc.add_heading('Strategy Synthesis', level=1)

    doc.add_heading('Business Issues', level=2)
    doc.add_paragraph(issue_text)

    doc.add_heading('Frameworks Applied', level=2)
    for fw in frameworks:
        doc.add_paragraph(fw, style='List Bullet')

    doc.add_heading('Vision Options', level=2)
    for i, v in enumerate(visions, start=1):
        p = doc.add_paragraph(f"{i}. {v}")
        if i - 1 == chosen_idx and p.runs:
            p.runs[0].bold = True

    doc.add_heading('Mission Statement', level=2)
    doc.add_paragraph(mission)

    doc.add_heading('Strategic Goals', level=2)
    for g in goals:
        doc.add_paragraph(g, style='List Number')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_selected_docx(selected_vision: str, mission: str, goals: List[str]) -> bytes:
    doc = Document()
    doc.add_heading('Selected Strategy', level=1)

    doc.add_heading('Vision', level=2)
    doc.add_paragraph(selected_vision)

    doc.add_heading('Mission', level=2)
    doc.add_paragraph(mission)

    doc.add_heading('Goals', level=2)
    for g in goals:
        doc.add_paragraph(g, style='List Number')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_multi_framework_docx(issue_text: str, frameworks_data: Dict[str, Dict[str, object]]) -> bytes:
    doc = Document()
    doc.add_heading('Multi-Framework Strategy Analysis', level=1)

    doc.add_heading('Business Issues / Context', level=2)
    doc.add_paragraph(issue_text)

    for fw, data in frameworks_data.items():
        doc.add_heading(fw, level=2)
        notes = (data.get("notes") or "").strip()
        if notes:
            doc.add_heading('Framework Notes', level=3)
            doc.add_paragraph(notes)
        visions = data.get("visions") or []
        if visions:
            doc.add_heading('Vision Options', level=3)
            selected_idx = data.get("selected_idx") or 0
            for i, v in enumerate(visions, start=1):
                p = doc.add_paragraph(f"{i}. {v}")
                if i - 1 == selected_idx and p.runs:
                    p.runs[0].bold = True
        mission = data.get("mission") or ""
        goals = data.get("goals") or []
        if mission:
            doc.add_heading('Mission Statement', level=3)
            doc.add_paragraph(str(mission))
        if goals:
            doc.add_heading('Strategic Goals', level=3)
            for g in goals:
                doc.add_paragraph(str(g), style='List Number')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def extract_docx_text(file) -> str:
    """Extract text from a .docx file, including paragraphs and table cells."""
    try:
        doc = Document(file)
    except Exception:
        return ""
    chunks: List[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            chunks.append(text)
    for table in getattr(doc, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    chunks.append(text)
    return "\n".join(chunks)


st.set_page_config(page_title="Vision & Strategy Synthesizer", page_icon="🧠", layout="wide")

st.title("🧠 Vision & Strategy Synthesizer")

with st.sidebar:
    st.header("Model Settings")
    host_val = st.text_input("Ollama Host", value=DEFAULT_OLLAMA_HOST, key="ollama_host")
    model_val = st.text_input("Model", value=DEFAULT_OLLAMA_MODEL, key="ollama_model")
    temperature = st.slider("Temperature", 0.0, 1.5, 0.7, 0.1)

st.markdown("Enter your business issues, optionally upload a Word document, select frameworks, and generate strategic options.")

issue_text = st.text_area("Business issues / context", height=180, placeholder="Describe the current situation, constraints, and opportunities...")

uploaded = st.file_uploader("Upload Word document (.docx)", type=["docx"], accept_multiple_files=False)
if uploaded is not None:
    uploaded_text = extract_docx_text(uploaded)
    st.session_state["uploaded_issue_text"] = uploaded_text
    with st.expander("Preview uploaded document", expanded=False):
        if uploaded_text:
            st.write(uploaded_text[:4000])
        else:
            st.write("(No extractable text found)")

include_uploaded_default = bool(st.session_state.get("uploaded_issue_text"))
include_uploaded = st.checkbox("Include uploaded document text", value=include_uploaded_default, key="include_uploaded") if include_uploaded_default else False

cols = st.columns(3)
selected_frameworks: List[str] = []
for i, fw in enumerate(FRAMEWORKS):
    col = cols[i % 3]
    with col:
        if st.checkbox(fw, value=False, key=f"fw_{i}"):
            selected_frameworks.append(fw)

uploaded_text_state = st.session_state.get("uploaded_issue_text", "")
combined_issue_text = (issue_text or "")
if include_uploaded and uploaded_text_state:
    combined_issue_text = (combined_issue_text + ("\n\n" if combined_issue_text else "") + uploaded_text_state)

# Framework template step
st.subheader("Framework templates (optional)")
if "framework_notes" not in st.session_state:
    st.session_state["framework_notes"] = {}
for fw in selected_frameworks:
    with st.expander(f"{fw} template", expanded=False):
        st.caption(FRAMEWORK_DESCRIPTIONS.get(fw, ""))
        st.write(FRAMEWORK_TIPS.get(fw, ""))
        st.session_state["framework_notes"][fw] = st.text_area(f"Notes for {fw}", key=f"notes_{fw}", height=120)

# Existing single combined generation path
st.markdown("---")
st.subheader("Combined-framework synthesis")
disable_generate = (not combined_issue_text) or (not selected_frameworks)
if st.button("Generate 3 one-sentence visions", type="primary", disabled=disable_generate):
    with st.spinner("Generating vision statements..."):
        try:
            prompt = build_vision_prompt(combined_issue_text, selected_frameworks)
            resp = ollama_generate(
                host=st.session_state.get("ollama_host", DEFAULT_OLLAMA_HOST),
                model=st.session_state.get("ollama_model", DEFAULT_OLLAMA_MODEL),
                system_prompt="You are a world-class strategy advisor.",
                user_prompt=prompt,
                temperature=temperature,
            )
            visions = parse_numbered_list(resp)
            if len(visions) < 3:
                visions = [ln.strip("-• ") for ln in resp.split("\n") if ln.strip()]
            visions = [first_sentence(v) for v in visions][:3]
            st.session_state["visions"] = visions
            st.session_state["issue_text"] = combined_issue_text
            st.session_state["frameworks"] = selected_frameworks
            st.session_state["model_settings"] = {
                "ollama_host": st.session_state.get("ollama_host", DEFAULT_OLLAMA_HOST),
                "ollama_model": st.session_state.get("ollama_model", DEFAULT_OLLAMA_MODEL),
                "temperature": temperature,
            }
            st.success(f"Generated {len(visions)} vision options.")
        except Exception as e:
            st.error(f"Failed to generate visions: {e}")

visions: List[str] = st.session_state.get("visions", [])
if visions:
    st.subheader("Vision Options")
    selected_idx = st.radio(
        "Select a preferred vision",
        options=list(range(len(visions))),
        format_func=lambda i: f"{i+1}. {visions[i]}",
        index=0,
        horizontal=False,
        key="vision_idx",
    )
    if st.button("Develop mission and goals", type="secondary"):
        with st.spinner("Drafting mission and goals..."):
            try:
                prompt = build_mission_goals_prompt(
                    selected_vision=visions[selected_idx],
                    issue_text=st.session_state.get("issue_text", ""),
                    selected_frameworks=st.session_state.get("frameworks", []),
                )
                resp = ollama_generate(
                    host=st.session_state.get("ollama_host", DEFAULT_OLLAMA_HOST),
                    model=st.session_state.get("ollama_model", DEFAULT_OLLAMA_MODEL),
                    system_prompt="You are a world-class strategy advisor.",
                    user_prompt=prompt,
                    temperature=temperature,
                )
                mission_raw = ""
                goals: List[str] = []
                for line in resp.splitlines():
                    stripped = line.strip()
                    if stripped.lower().startswith("mission:"):
                        mission_raw = stripped.split(":", 1)[1].strip()
                    elif stripped and stripped[0].isdigit() and stripped[1:2] == ".":
                        goals.append(stripped.split(".", 1)[1].strip())
                if not mission_raw:
                    mission_raw = resp.split("\n", 1)[0].strip()
                mission = clean_mission_text(resp, mission_raw)
                if len(mission.split()) < 3 or is_meta_mission(mission):
                    mission_regen = generate_mission_only(
                        host=st.session_state.get("ollama_host", DEFAULT_OLLAMA_HOST),
                        model=st.session_state.get("ollama_model", DEFAULT_OLLAMA_MODEL),
                        vision=visions[selected_idx],
                        issue_text=st.session_state.get("issue_text", ""),
                        temperature=temperature,
                    )
                    mission_regen = clean_mission_text(mission_regen, mission_regen)
                    if len(mission_regen.split()) >= 3 and not is_meta_mission(mission_regen):
                        mission = mission_regen
                    else:
                        synthesized = synthesize_mission_from_vision(visions[selected_idx])
                        if synthesized:
                            mission = synthesized
                if len(goals) < 5:
                    goals = parse_numbered_list(resp)[:5]
                st.session_state["mission"] = mission
                st.session_state["goals"] = goals
                st.success("Mission and goals drafted.")
            except Exception as e:
                st.error(f"Failed to generate mission/goals: {e}")

mission = st.session_state.get("mission")
goals = st.session_state.get("goals", [])
if mission or goals:
    st.subheader("Results")
    if mission:
        st.markdown(f"**Mission:** {mission}")
    if goals:
        st.markdown("**Goals:**")
        for i, g in enumerate(goals, start=1):
            st.markdown(f"{i}. {g}")

    selected_idx = st.session_state.get("vision_idx", 0)
    visions_list: List[str] = st.session_state.get("visions", [])
    selected_vision = visions_list[selected_idx] if visions_list else ""
    selected_content = export_selected_docx(
        selected_vision=selected_vision,
        mission=mission or "",
        goals=goals,
    )
    st.download_button(
        label="Download selected strategy.docx",
        data=selected_content,
        file_name="selected_strategy.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    full_content = export_docx(
        issue_text=st.session_state.get("issue_text", ""),
        frameworks=st.session_state.get("frameworks", []),
        visions=st.session_state.get("visions", []),
        chosen_idx=selected_idx,
        mission=mission or "",
        goals=goals,
    )
    st.download_button(
        label="Download full analysis.docx",
        data=full_content,
        file_name="analysis.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# Per-framework generation
st.markdown("---")
st.subheader("Per-framework analysis")
if "visions_by_fw" not in st.session_state:
    st.session_state["visions_by_fw"] = {}
if "vision_idx_by_fw" not in st.session_state:
    st.session_state["vision_idx_by_fw"] = {}
if "mission_goals_by_fw" not in st.session_state:
    st.session_state["mission_goals_by_fw"] = {}

pf_disabled = (not combined_issue_text) or (not selected_frameworks)
if st.button("Generate 3 visions per selected framework", disabled=pf_disabled):
    try:
        for fw in selected_frameworks:
            notes = (st.session_state["framework_notes"].get(fw) if "framework_notes" in st.session_state else "") or ""
            prompt = build_framework_vision_prompt(combined_issue_text, fw, notes)
            resp = ollama_generate(
                host=st.session_state.get("ollama_host", DEFAULT_OLLAMA_HOST),
                model=st.session_state.get("ollama_model", DEFAULT_OLLAMA_MODEL),
                system_prompt="You are a world-class strategy advisor.",
                user_prompt=prompt,
                temperature=temperature,
            )
            visions_fw = parse_numbered_list(resp)
            if len(visions_fw) < 3:
                visions_fw = [ln.strip("-• ") for ln in resp.split("\n") if ln.strip()]
            visions_fw = [first_sentence(v) for v in visions_fw][:3]
            st.session_state["visions_by_fw"][fw] = visions_fw
            if fw not in st.session_state["vision_idx_by_fw"]:
                st.session_state["vision_idx_by_fw"][fw] = 0
        st.success("Generated per-framework visions.")
    except Exception as e:
        st.error(f"Failed to generate per-framework visions: {e}")

# Selection and mission/goals per framework
for fw in selected_frameworks:
    visions_fw = st.session_state["visions_by_fw"].get(fw, [])
    if not visions_fw:
        continue
    st.markdown(f"**{fw} visions**")
    key_idx = f"vision_idx_{fw}"
    st.session_state["vision_idx_by_fw"][fw] = st.radio(
        f"Select a vision for {fw}",
        options=list(range(len(visions_fw))),
        format_func=lambda i, vf=visions_fw: f"{i+1}. {vf[i]}",
        index=st.session_state["vision_idx_by_fw"].get(fw, 0),
        horizontal=False,
        key=key_idx,
    )

if st.button("Develop missions and goals per framework"):
    try:
        for fw in selected_frameworks:
            visions_fw = st.session_state["visions_by_fw"].get(fw, [])
            if not visions_fw:
                continue
            sel_idx = st.session_state["vision_idx_by_fw"].get(fw, 0)
            sel_idx = min(max(sel_idx, 0), len(visions_fw) - 1)
            sel_vision = visions_fw[sel_idx]
            prompt = build_mission_goals_prompt(
                selected_vision=sel_vision,
                issue_text=combined_issue_text,
                selected_frameworks=[fw],
            )
            resp = ollama_generate(
                host=st.session_state.get("ollama_host", DEFAULT_OLLAMA_HOST),
                model=st.session_state.get("ollama_model", DEFAULT_OLLAMA_MODEL),
                system_prompt="You are a world-class strategy advisor.",
                user_prompt=prompt,
                temperature=temperature,
            )
            mission_raw = ""
            goals: List[str] = []
            for line in resp.splitlines():
                stripped = line.strip()
                if stripped.lower().startswith("mission:"):
                    mission_raw = stripped.split(":", 1)[1].strip()
                elif stripped and stripped[0].isdigit() and stripped[1:2] == ".":
                    goals.append(stripped.split(".", 1)[1].strip())
            if not mission_raw:
                mission_raw = resp.split("\n", 1)[0].strip()
            mission = clean_mission_text(resp, mission_raw)
            if len(mission.split()) < 3 or is_meta_mission(mission):
                mission_regen = generate_mission_only(
                    host=st.session_state.get("ollama_host", DEFAULT_OLLAMA_HOST),
                    model=st.session_state.get("ollama_model", DEFAULT_OLLAMA_MODEL),
                    vision=sel_vision,
                    issue_text=combined_issue_text,
                    temperature=temperature,
                )
                mission_regen = clean_mission_text(mission_regen, mission_regen)
                if len(mission_regen.split()) >= 3 and not is_meta_mission(mission_regen):
                    mission = mission_regen
                else:
                    synthesized = synthesize_mission_from_vision(sel_vision)
                    if synthesized:
                        mission = synthesized
            if len(goals) < 5:
                goals = parse_numbered_list(resp)[:5]
            st.session_state["mission_goals_by_fw"][fw] = {
                "notes": st.session_state.get("framework_notes", {}).get(fw, ""),
                "visions": visions_fw,
                "selected_idx": sel_idx,
                "mission": mission,
                "goals": goals,
            }
        st.success("Generated missions and goals per framework.")
        # Immediately render results inline
        st.subheader("Per-framework results")
        for fw in selected_frameworks:
            data = st.session_state["mission_goals_by_fw"].get(fw)
            if not data:
                continue
            st.markdown(f"### {fw}")
            visions_fw = data.get("visions", [])
            sel_idx = data.get("selected_idx", 0)
            if visions_fw:
                st.markdown(f"Selected vision: {visions_fw[sel_idx] if 0 <= sel_idx < len(visions_fw) else ''}")
            mission_fw = data.get("mission", "")
            goals_fw = data.get("goals", [])
            if mission_fw:
                st.markdown(f"**Mission:** {mission_fw}")
            if goals_fw:
                st.markdown("**Goals:**")
                for i, g in enumerate(goals_fw, start=1):
                    st.markdown(f"{i}. {g}")
    except Exception as e:
        st.error(f"Failed to generate per-framework mission/goals: {e}")

# Display per-framework results before download
if st.session_state.get("mission_goals_by_fw"):
    st.subheader("Per-framework results")
    for fw in selected_frameworks:
        data = st.session_state["mission_goals_by_fw"].get(fw)
        if not data:
            continue
        st.markdown(f"### {fw}")
        visions_fw = data.get("visions", [])
        sel_idx = data.get("selected_idx", 0)
        if visions_fw:
            st.markdown(f"Selected vision: {visions_fw[sel_idx] if 0 <= sel_idx < len(visions_fw) else ''}")
        mission_fw = data.get("mission", "")
        goals_fw = data.get("goals", [])
        if mission_fw:
            st.markdown(f"**Mission:** {mission_fw}")
        if goals_fw:
            st.markdown("**Goals:**")
            for i, g in enumerate(goals_fw, start=1):
                st.markdown(f"{i}. {g}")

    content = export_multi_framework_docx(
        issue_text=combined_issue_text,
        frameworks_data=st.session_state["mission_goals_by_fw"],
    )
    st.download_button(
        label="Download multi-framework analysis.docx",
        data=content,
        file_name="multi_framework_analysis.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )